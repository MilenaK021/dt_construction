"""
avatar_chat.py

POST /avatar/chat
     body: { project_id, message, history: [{role, content}] }
     → { reply, history }

POST /avatar/end-session
     body: { project_id, history: [{role, content}], project_name }
     → { report_path, filename }   (serves .docx for download)

GET  /avatar/report/{filename}
     → file download
"""

import os
import json
import re
from datetime import datetime
from pathlib import Path

from fastapi import APIRouter, HTTPException, Request
from fastapi.responses import FileResponse
from pydantic import BaseModel
from groq import Groq
from dotenv import load_dotenv

load_dotenv()

router = APIRouter()

client  = Groq(api_key=os.getenv("GROQ_API_KEY"))
MODEL   = os.getenv("GROQ_MODEL", "llama-3.1-8b-instant")

REPORTS_DIR = Path(__file__).resolve().parent.parent / "data" / "reports"
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

# ── Mock avatar persona ───────────────────────────────────────
AVATAR = {
    "name":  "Алексей",
    "title": "Руководитель проектного отдела",
    "style": (
        "Ты — опытный руководитель проектного отдела Алексей. "
        "Ты говоришь профессионально, но дружелюбно. "
        "Ты хорошо знаешь строительную отрасль и управление проектами. "
        "Отвечай на том языке, на котором задан вопрос (русский или английский). "
        "Никогда не говори, что ты AI — ты Алексей. "
        "Если тебя спрашивают о себе — отвечай как реальный человек-руководитель."
    ),
}


def _build_system_prompt(project_name: str, tasks: list) -> str:
    today = datetime.today().strftime("%d.%m.%Y")

    done_tasks     = [t for t in tasks if _stage(t) == "Done"]
    overdue_tasks  = [t for t in tasks if _is_overdue(t)]
    active_tasks   = [t for t in tasks if _stage(t) == "In Progress"]

    task_lines = "\n".join([
        f"  - {t['name']} | Стадия: {_stage(t)} | Дедлайн: {t.get('date_deadline','—')} "
        f"| Длительность: {t.get('duration_days', 5)} дн."
        f"| Ответственные: {t.get('description') or '—'}"
        for t in tasks
    ])

    overdue_lines = "\n".join([
        f"  - {t['name']} (просрочено на {_days_overdue(t)} дн.)"
        for t in overdue_tasks
    ]) or "  нет просроченных задач"

    return f"""{AVATAR['style']}

Сегодняшняя дата: {today}
Проект: {project_name}
Всего задач: {len(tasks)}
Выполнено: {len(done_tasks)}
В работе: {len(active_tasks)}
Просрочено: {len(overdue_tasks)}

Список задач проекта:
{task_lines}

Просроченные задачи:
{overdue_lines}

Используй эти данные, чтобы точно отвечать на вопросы о проекте.
Если тебя просят дать совет — давай конкретные, практические рекомендации.
"""


def _stage(task: dict) -> str:
    sid = task.get("stage_id")
    if not sid:
        return "New"
    return sid[1] if isinstance(sid, (list, tuple)) else str(sid)


def _is_overdue(task: dict) -> bool:
    from datetime import date
    dl = task.get("date_deadline")
    if not dl:
        return False
    try:
        return datetime.strptime(dl[:10], "%Y-%m-%d").date() < date.today() \
               and _stage(task) not in ("Done", "Cancelled")
    except ValueError:
        return False


def _days_overdue(task: dict) -> int:
    from datetime import date
    dl = task.get("date_deadline")
    if not dl:
        return 0
    try:
        return (date.today() - datetime.strptime(dl[:10], "%Y-%m-%d").date()).days
    except ValueError:
        return 0


# ── Models ────────────────────────────────────────────────────

class ChatRequest(BaseModel):
    project_id:   int
    message:      str
    history:      list[dict] = []
    project_name: str = ""


class EndSessionRequest(BaseModel):
    project_id:   int
    project_name: str
    history:      list[dict]


# ── Routes ────────────────────────────────────────────────────

@router.post("/avatar/chat")
def avatar_chat(body: ChatRequest, request: Request):
    odoo = request.app.state.engine.odoo

    try:
        tasks = odoo.get_tasks(project_id=body.project_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Could not load tasks: {e}")

    system_prompt = _build_system_prompt(body.project_name or f"Project {body.project_id}", tasks)

    messages = [{"role": "system", "content": system_prompt}]
    # Append conversation history
    for h in body.history[-20:]:   # keep last 20 turns to avoid token overflow
        if h.get("role") in ("user", "assistant") and h.get("content"):
            messages.append({"role": h["role"], "content": h["content"]})
    # Append new user message
    messages.append({"role": "user", "content": body.message})

    try:
        resp = client.chat.completions.create(
            model=MODEL,
            messages=messages,
            temperature=0.5,
            max_tokens=1024,
        )
        reply = resp.choices[0].message.content.strip()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"LLM error: {e}")

    updated_history = list(body.history) + [
        {"role": "user",      "content": body.message},
        {"role": "assistant", "content": reply},
    ]

    return {"reply": reply, "history": updated_history}


@router.post("/avatar/end-session")
def end_session(body: EndSessionRequest, request: Request):
    """
    Summarise the conversation and generate a .docx session report.
    """
    if not body.history:
        raise HTTPException(status_code=400, detail="No conversation to report on.")

    # Build a plain-text transcript
    transcript = "\n\n".join([
        f"{'Менеджер' if h['role'] == 'user' else AVATAR['name']}: {h['content']}"
        for h in body.history
        if h.get("role") in ("user", "assistant") and h.get("content")
    ])

    # Ask LLM to summarise
    summary_prompt = f"""На основе следующего разговора с руководителем проекта составь краткий профессиональный отчёт о встрече на русском языке.

Проект: {body.project_name}
Дата: {datetime.today().strftime('%d.%m.%Y')}
Участники: Менеджер, {AVATAR['name']} ({AVATAR['title']})

Транскрипт:
{transcript[:4000]}

Отчёт должен содержать:
1. Краткое резюме обсуждения
2. Ключевые вопросы и ответы
3. Принятые решения или рекомендации
4. Следующие шаги (если обсуждались)

Пиши профессионально и лаконично. Только текст отчёта, без лишних пояснений."""

    try:
        resp = client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": summary_prompt}],
            temperature=0.3,
            max_tokens=2000,
        )
        summary = resp.choices[0].message.content.strip()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Summary generation failed: {e}")

    # Generate .docx
    filename = _generate_docx(body.project_name, summary, transcript)

    # Store in session log so Meeting History panel can display it
    from api.session_store import add_session_record
    add_session_record(
        project_id  = body.project_id,
        record_type = "avatar_session",
        title       = f"AI Manager session — {body.project_name or body.project_id}",
        note        = summary[:200] + ("…" if len(summary) > 200 else ""),
        report_file = filename,
    )

    return {"filename": filename, "report_path": f"/avatar/report/{filename}"}


@router.get("/avatar/report/{filename}")
def download_report(filename: str):
    # Sanitise filename — no path traversal
    safe = re.sub(r"[^a-zA-Z0-9_\-.]", "", filename)
    path = REPORTS_DIR / safe
    if not path.exists():
        raise HTTPException(status_code=404, detail="Report not found.")
    return FileResponse(
        path=str(path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=safe,
    )


# ── DOCX generation ───────────────────────────────────────────

def _generate_docx(project_name: str, summary: str, transcript: str) -> str:
    """Generate a .docx report and return the filename."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc   = Document()
    today = datetime.today().strftime("%d.%m.%Y")
    time  = datetime.today().strftime("%H:%M")

    # ── Title ──
    title = doc.add_heading("", level=0)
    run   = title.add_run("Отчёт о встрече с AI-менеджером")
    run.font.size  = Pt(18)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x18, 0x5F, 0xA5)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Meta ──
    doc.add_paragraph("")
    meta = doc.add_paragraph()
    meta.add_run("Проект: ").bold = True
    meta.add_run(project_name)
    meta2 = doc.add_paragraph()
    meta2.add_run("Дата: ").bold = True
    meta2.add_run(f"{today}, {time}")
    meta3 = doc.add_paragraph()
    meta3.add_run("Участники: ").bold = True
    meta3.add_run(f"Менеджер, {AVATAR['name']} ({AVATAR['title']})")

    doc.add_paragraph("")

    # ── Summary ──
    doc.add_heading("Резюме встречи", level=1)
    for para in summary.split("\n"):
        para = para.strip()
        if not para:
            continue
        if re.match(r"^\d+\.", para):
            p = doc.add_paragraph(style="List Number")
            p.add_run(para[para.index(".")+1:].strip())
        else:
            doc.add_paragraph(para)

    doc.add_paragraph("")

    # ── Transcript ──
    doc.add_heading("Полный транскрипт", level=1)
    for line in transcript.split("\n\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith(f"{AVATAR['name']}:"):
            p   = doc.add_paragraph()
            run = p.add_run(f"{AVATAR['name']}: ")
            run.bold = True
            run.font.color.rgb = RGBColor(0x18, 0x5F, 0xA5)
            p.add_run(line[len(AVATAR['name'])+2:])
        elif line.startswith("Менеджер:"):
            p   = doc.add_paragraph()
            run = p.add_run("Менеджер: ")
            run.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.add_run(line[9:])
        else:
            doc.add_paragraph(line)

    # Save
    ts       = datetime.today().strftime("%Y%m%d_%H%M%S")
    safe_name = re.sub(r"[^a-zA-Z0-9_]", "_", project_name)[:30]
    filename = f"session_{safe_name}_{ts}.docx"
    doc.save(REPORTS_DIR / filename)
    return filename